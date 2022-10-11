from PdfEditor import PdfEditor
import tkinter as tk
from tkinter import filedialog
from tkinter import font as tkfont
import os
from PyPDF2 import PdfFileReader, PdfFileWriter,PdfFileMerger
from docx import Document
from docx.shared import Inches
import sys
import comtypes.client

class Application(tk.Tk):
    def __init__(self):
        tk.Tk.__init__(self)
        # Main window is not resizable for now 
        self.resizable(0, 0)
        self.geometry("350x280")
        self.title("PDFEdit")
        self.create_widgets()

    # Creating widgets for the starting screen 
    def create_widgets(self):
        self.merge_btn = tk.Button(self, font =("Arial", "12"), text = "Merge PDF", padx = 30, pady = 30, command = self.merge, fg = '#ffffff', bg = '#ff0000')
        self.to_docx_btn = tk.Button(self,font =("Arial", "12"), text = "PDF to .docx", padx = 25, pady = 30, fg = '#ffffff',  bg = '#ff0000', command = self.pdf_to_docx)
        self.edit_btn = tk.Button(self, font =("Arial", "12"), text = "Edit PDF", padx = 39, pady = 30, fg = '#ffffff',  bg = '#ff0000', command = self.edit_pdf)
        self.to_pdf_btn = tk.Button(self, font =("Arial", "12"), text = ".docx to PDF", padx = 25, pady = 30, fg = '#ffffff', bg = '#ff0000', command = self.docx_to_pdf)
        # Positioning the buttons 
        self.merge_btn.grid(row = 2, column = 2,  padx = 10, pady = 20)
        self.to_docx_btn.grid(row = 2, column = 4,  padx = 10, pady = 20)
        self.edit_btn.grid(row = 4, column = 2,  padx = 10, pady = 20)
        self.to_pdf_btn.grid(row = 4, column = 4,  padx = 10, pady = 20)


    def merge(self):
        top = tk.Toplevel()
        top.title("Edit your Pdf")
        top.iconbitmap('pdficon.ico')
        top.geometry("400x400")
        f_label = tk.Label(top, font =("Arial", "12"), text = "Merged file name:", pady = 20)
        file = tk.Entry(top, width = 35, borderwidth = 5)
        button = tk.Button(top, font =("Arial", "12"), text = "Merge PDF", padx = 30, pady = 20, 
                            fg = '#ffffff', bg = '#ff0000', command = lambda: self.merge_wrapper(file.get()))
        f_label.pack()
        file.pack()
        button.pack()


    # This function opens 3 separate file dialogs to find the combined PDF and location 
    # to be saved
    # Merged file is opened afterwards 
    def merge_wrapper(self, output):
        pdf1 = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select first PDF")
        pdf2 = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select second PDF")
        location = tk.filedialog.askdirectory(initialdir = os.getcwd(), title = "Select a folder to be stored")
        output=PdfFileMerger()
        output.append(pdf1)
        output.append(pdf2)
        output.write(open(location, 'w+'))
        os.startfile(combined)

    def pdf_to_docx(self):
        document = Document()
        document.add_heading('Document Title', 0)
        path = filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select PDF")
        text=""
        pdf_file = open(path, 'rb')
        read_pdf = PdfFileReader(pdf_file)
        c = read_pdf.numPages
        print('gunga')
        for i in range(c):
          page = read_pdf.getPage(i)
          text+=(page.extractText())
        document.add_paragraph(text)
        document.add_page_break()
        document.save(path + '.docx')

    def docx_to_pdf(self):
        path = filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select Word Document")
        wdFormatPDF = 17
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(path)
        doc.SaveAs(path + '.pdf', FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

    # Edit_pdf creates a new window to create a new_pdf from a user specified page range
    def edit_pdf(self):
        top = tk.Toplevel()
        top.title("Edit your Pdf")
        top.iconbitmap('pdficon.ico')
        top.geometry("400x400")
        label = tk.Label(top, font =("Arial", "12"), text = "Page range or pages (e.g 1-2 or 1,2,3)",pady = 20)
        label3 = tk.Label(top, font =("Arial", "12"), text = "File name", pady = 20)
        pages = tk.Entry(top, width = 35, borderwidth = 5)
        file = tk.Entry(top, width = 35, borderwidth = 5)
        button = tk.Button(top, font =("Arial", "12"), text = "Create PDF", padx = 30, pady = 20, 
                            fg = '#ffffff', bg = '#ff0000', command = lambda: 
                            self.create_split_pdf(file.get()+".pdf", pages.get()))
        label.pack()
        pages.pack()
        label3.pack()
        file.pack()
        button.pack()


    # A wrapper function for create_split in pdf_edit.py
    # Opens newly created pdf and prompts user to select pdf and folder to store
    def create_split_pdf(self, output_name, pages): 
        file = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select PDF")
        folder = tk.filedialog.askdirectory(initialdir = os.getcwd(), title = "Select Folder To Store")
        new_file = edit.create_split(file, output_name, folder, pages)
        os.startfile(new_file)


    def start_file(self, new_file):
        os.startfile(new_file)

if __name__ == '__main__':
    app = Application()
    app.mainloop()
